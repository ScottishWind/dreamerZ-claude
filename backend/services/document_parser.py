"""Document parser — extract text and structure from PDF/DOCX uploads.

Used by the course-creator pipeline to convert raw admin uploads into
clean, structured text ready for LLM analysis.
"""

import io
import logging
from typing import Optional

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None  # type: ignore

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore


MAX_RAW_TEXT_CHARS = 200_000  # Hard cap to protect downstream LLM prompts


class UnsupportedFormatError(ValueError):
    """Raised when a file type is not recognised by the parser."""


def _truncate(text: str) -> str:
    if len(text) <= MAX_RAW_TEXT_CHARS:
        return text
    logging.warning(
        "Parsed text truncated from %d to %d chars",
        len(text),
        MAX_RAW_TEXT_CHARS,
    )
    return text[:MAX_RAW_TEXT_CHARS]


def parse_pdf_bytes(data: bytes) -> dict:
    """Extract text and naive heading structure from a PDF byte blob."""
    if PdfReader is None:
        raise RuntimeError(
            "pypdf is not installed. Add 'pypdf' to requirements.txt."
        )

    reader = PdfReader(io.BytesIO(data))
    text_parts: list[str] = []
    structure: list[dict] = []

    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover
            logging.warning("Failed to extract page %d: %s", page_idx, exc)
            page_text = ""

        text_parts.append(page_text)

        # Cheap heading detection: non-empty short lines followed by blank line
        for line in page_text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if 4 <= len(stripped) <= 90 and stripped == stripped.strip(":.,-"):
                # Heuristic: looks like a heading if it's short and title-ish
                if stripped[0].isupper() and sum(c.isupper() for c in stripped) >= 2:
                    structure.append(
                        {"text": stripped, "type": "heading", "page": page_idx}
                    )

    full_text = "\n".join(text_parts).strip()
    return {
        "raw_text": _truncate(full_text),
        "structure": structure[:200],
        "page_count": len(reader.pages),
        "char_count": len(full_text),
    }


def parse_docx_bytes(data: bytes) -> dict:
    """Extract text and heading structure from a DOCX byte blob."""
    if docx is None:
        raise RuntimeError(
            "python-docx is not installed. Add 'python-docx' to requirements.txt."
        )

    document = docx.Document(io.BytesIO(data))
    text_parts: list[str] = []
    structure: list[dict] = []

    for para in document.paragraphs:
        para_text = (para.text or "").strip()
        if not para_text:
            continue
        text_parts.append(para_text)

        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading"):
            level = 1
            # Extract heading level from "Heading 1", "Heading 2", ...
            for token in style_name.split():
                if token.isdigit():
                    level = int(token)
                    break
            structure.append({"text": para_text, "type": "heading", "level": level})

    # Extract table contents as plain text so they contribute to the context
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n".join(text_parts).strip()
    return {
        "raw_text": _truncate(full_text),
        "structure": structure[:200],
        "paragraph_count": len(document.paragraphs),
        "char_count": len(full_text),
    }


def parse_document(
    data: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> dict:
    """Dispatch parsing based on file extension / mime type.

    Returns a dict with keys: raw_text, structure, ...metadata.
    Raises UnsupportedFormatError for unknown formats.
    """
    name = (filename or "").lower()
    mime = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in mime:
        return parse_pdf_bytes(data)

    if name.endswith(".docx") or "wordprocessingml" in mime:
        return parse_docx_bytes(data)

    # Plain text fallback (useful for quick admin pastes via .txt/.md)
    if name.endswith((".txt", ".md")) or mime.startswith("text/"):
        text = data.decode("utf-8", errors="replace")
        return {
            "raw_text": _truncate(text),
            "structure": [],
            "char_count": len(text),
        }

    raise UnsupportedFormatError(
        f"Unsupported document type: '{filename}'. "
        "Accepted formats: PDF, DOCX, TXT, MD."
    )
